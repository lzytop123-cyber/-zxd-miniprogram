"""知行岛 AI 学习助手 — DeepSeek 调用 + 知识库/上下文组装。

设计：轻量方案。把「可编辑知识文档 + 当前用户学习报告数据」拼入 system 提示词，
不引入向量检索。DeepSeek 为 OpenAI 兼容接口，使用 httpx 同步调用（与 report 等同步路由一致）。
"""

import logging
from datetime import date, timedelta
from functools import lru_cache
from io import BytesIO
from pathlib import Path

import httpx
from sqlalchemy import func, select
from sqlalchemy.orm import Session

from app.core.config import settings
from app.models import StudyStat, User

logger = logging.getLogger(__name__)

KNOWLEDGE_PATH = Path(__file__).resolve().parent.parent / "knowledge" / "zhixingdao_kb.md"
KNOWLEDGE_ALLOWED_EXTENSIONS = {".md", ".markdown", ".txt", ".docx"}
KNOWLEDGE_MAX_BYTES = 2_000_000

# 历史消息限制，控制 token 成本
MAX_HISTORY_MESSAGES = 12
MAX_MESSAGE_CHARS = 1000

ASSISTANT_NOT_CONFIGURED_REPLY = (
    "AI 学习助手还没有配置好（缺少 DeepSeek API Key）。"
    "请在后端 backend/.env 填写 DEEPSEEK_API_KEY 后重启服务即可使用。"
)


@lru_cache(maxsize=1)
def load_knowledge() -> str:
    """读取知识文档（带进程内缓存）。"""
    try:
        return KNOWLEDGE_PATH.read_text(encoding="utf-8").strip()
    except FileNotFoundError:
        logger.warning("知识文档不存在: %s", KNOWLEDGE_PATH)
        return ""


def save_knowledge(content: str) -> None:
    """保存知识文档并清除缓存。"""
    KNOWLEDGE_PATH.parent.mkdir(parents=True, exist_ok=True)
    KNOWLEDGE_PATH.write_text(content.strip() + "\n", encoding="utf-8")
    load_knowledge.cache_clear()


def decode_text_upload(raw: bytes) -> str:
    """解析上传的文本文档（UTF-8 / GBK）。"""
    if not raw:
        raise ValueError("文档内容为空")
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "gbk"):
        try:
            text = raw.decode(encoding).strip()
            if text:
                return text
        except UnicodeDecodeError:
            continue
    raise ValueError("无法识别文件编码，请使用 UTF-8 或 GBK 的 .md / .txt 文件")


def decode_docx_upload(raw: bytes) -> str:
    """从 Word (.docx) 提取纯文本。"""
    try:
        from docx import Document
    except ImportError as e:
        raise ValueError("服务器未安装 Word 解析依赖，请联系管理员") from e

    try:
        doc = Document(BytesIO(raw))
    except Exception as e:
        raise ValueError("无法解析 Word 文档，请使用 .docx（不支持旧版 .doc）") from e

    lines: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))

    content = "\n".join(lines).strip()
    if not content:
        raise ValueError("Word 文档中没有可提取的文字")
    return content


def parse_knowledge_upload(raw: bytes, *, ext: str) -> str:
    """按扩展名解析上传文档。"""
    suffix = (ext or "").lower()
    if suffix == ".docx":
        return decode_docx_upload(raw)
    if suffix in {".md", ".markdown", ".txt"}:
        return decode_text_upload(raw)
    raise ValueError("不支持的文件格式")


def apply_knowledge_upload(text: str, *, mode: str = "replace") -> str:
    """将上传文档合并进知识库内容。"""
    text = (text or "").strip()
    if not text:
        raise ValueError("文档内容为空")
    if mode == "append":
        existing = load_knowledge().strip()
        if existing:
            return f"{existing}\n\n---\n\n{text}"
    return text


def build_user_context(db: Session, user: User) -> str:
    """根据 StudyStat 拼出该用户的学习概况（口径与 report.summary 一致）。"""
    total_minutes = int(
        db.scalar(select(func.coalesce(func.sum(StudyStat.total_minutes), 0)).where(StudyStat.user_id == user.id))
        or 0
    )
    total_sessions = int(
        db.scalar(select(func.coalesce(func.sum(StudyStat.session_count), 0)).where(StudyStat.user_id == user.id))
        or 0
    )
    distinct_days = db.scalar(
        select(func.count(func.distinct(StudyStat.stat_date))).where(StudyStat.user_id == user.id)
    ) or 0
    daily_avg = total_minutes // max(distinct_days, 1)

    since = date.today() - timedelta(days=6)
    recent_minutes = int(
        db.scalar(
            select(func.coalesce(func.sum(StudyStat.total_minutes), 0)).where(
                StudyStat.user_id == user.id, StudyStat.stat_date >= since
            )
        )
        or 0
    )

    nickname = user.nickname or "同学"
    title = user.title or "小白"
    lines = [
        f"- 昵称：{nickname}（称号：{title}）",
        f"- 累计学习：{total_minutes // 60} 小时 {total_minutes % 60} 分，共 {total_sessions} 次，打卡 {distinct_days} 天",
        f"- 日均时长：约 {daily_avg} 分钟",
        f"- 近 7 日学习：{recent_minutes // 60} 小时 {recent_minutes % 60} 分",
    ]
    if total_minutes == 0:
        lines.append("- 该用户暂无学习记录，可鼓励 TA 开启第一次专注。")
    return "\n".join(lines)


def build_system_prompt(user_context: str, *, query: str = "") -> str:
    from app.services import knowledge_rag

    knowledge_rag.migrate_legacy_markdown_if_needed(load_knowledge)

    rag_context = ""
    if knowledge_rag.rag_available() and (query or "").strip():
        rag_context = knowledge_rag.retrieve_context(query)

    if rag_context:
        knowledge_block = rag_context
        knowledge_title = "门店知识库（检索到的相关片段）"
        knowledge_rule = (
            "门店相关问题只能依据下方【门店知识库】片段回答；片段中没有的，不要编造，"
            "应坦诚告知并建议「在小程序对应页面查看实际信息或联系店长」。"
        )
    else:
        knowledge_block = load_knowledge() or "（暂无知识库内容）"
        knowledge_title = "门店知识库"
        knowledge_rule = (
            "门店相关问题只能依据下方【门店知识库】回答；知识库中没有或标注「待补充」的，不要编造，"
            "应坦诚告知并建议「在小程序对应页面查看实际信息或联系店长」。"
        )

    return f"""你是「知行岛自习室」的 AI 学习助手，名字叫小岛。你的职责：
1. 解答门店相关问题（价格、营业时间、预约、入座、团购核销、会员积分等）；
2. 结合用户的学习数据，给出专注、时间管理、复习规划等学习建议；
3. 回答通用学习/答疑问题。

回答要求：
- 用中文，语气友好、简洁，多用分点；不要长篇大论。
- {knowledge_rule}
- 给学习建议时可引用下方【用户学习概况】，但不要泄露与对话无关的隐私。
- 不提供医疗、心理、法律等专业诊断；涉及健康问题建议咨询专业人士。

【{knowledge_title}】
{knowledge_block}

【用户学习概况】
{user_context}
"""


def _sanitize_history(messages: list[dict]) -> list[dict]:
    """裁剪历史消息：仅保留 user/assistant，限制条数与长度。"""
    cleaned: list[dict] = []
    for m in messages:
        role = m.get("role")
        content = (m.get("content") or "").strip()
        if role not in ("user", "assistant") or not content:
            continue
        cleaned.append({"role": role, "content": content[:MAX_MESSAGE_CHARS]})
    return cleaned[-MAX_HISTORY_MESSAGES:]


def chat(system_prompt: str, history: list[dict]) -> str:
    """调用 DeepSeek chat/completions，返回助手回复文本。"""
    if not settings.assistant_configured:
        return ASSISTANT_NOT_CONFIGURED_REPLY

    payload = {
        "model": settings.deepseek_model,
        "messages": [{"role": "system", "content": system_prompt}, *_sanitize_history(history)],
        "temperature": 0.7,
        "stream": False,
    }
    base = settings.deepseek_base_url.rstrip("/")
    headers = {
        "Authorization": f"Bearer {settings.deepseek_api_key}",
        "Content-Type": "application/json",
    }
    try:
        with httpx.Client(timeout=settings.deepseek_timeout_sec) as client:
            resp = client.post(f"{base}/chat/completions", headers=headers, json=payload)
        if resp.status_code >= 400:
            logger.warning("DeepSeek 返回错误 %s: %s", resp.status_code, resp.text[:500])
            return "AI 助手暂时不可用，请稍后再试～"
        data = resp.json()
        return (data["choices"][0]["message"]["content"] or "").strip() or "（助手没有返回内容）"
    except (httpx.HTTPError, KeyError, IndexError, ValueError) as exc:
        logger.warning("DeepSeek 调用失败: %s", exc)
        return "AI 助手暂时不可用，请稍后再试～"
